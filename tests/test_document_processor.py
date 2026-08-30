import pytest

from pathlib import Path

from docx import Document

from scripts.document_processor import (
    convert_docx,
    convert_txt,
    process_document,
)

def test_convert_txt_marks_short_uppercase_lines_as_headings(tmp_path: Path):
    source = tmp_path / "sample.txt"
    source.write_text(
        "INTRODUCTION\n\nThis is normal text.",
        encoding="utf-8",
    )

    result = convert_txt(source)

    assert "## Introduction" in result
    assert "This is normal text." in result


def test_convert_docx_preserves_heading_and_paragraph(tmp_path: Path):
    source = tmp_path / "sample.docx"

    doc = Document()
    doc.add_heading("Project Objectives", level=1)
    doc.add_paragraph("The project aims to improve adult learning participation.")
    doc.save(source)

    result = convert_docx(source)

    assert "# Project Objectives" in result
    assert "The project aims to improve adult learning participation." in result


def test_process_document_returns_none_for_unsupported_extension(tmp_path: Path):
    source = tmp_path / "sample.csv"
    source.write_text("a,b,c", encoding="utf-8")

    assert process_document(source) is None


def test_process_document_dispatches_txt(tmp_path: Path):
    source = tmp_path / "sample.txt"
    source.write_text("Simple text", encoding="utf-8")

    result = process_document(source)

    assert result == "Simple text"

@pytest.mark.filterwarnings(
    "ignore:builtin type SwigPyPacked has no __module__ attribute:DeprecationWarning"
)
@pytest.mark.filterwarnings(
    "ignore:builtin type SwigPyObject has no __module__ attribute:DeprecationWarning"
)
@pytest.mark.filterwarnings(
    "ignore:builtin type swigvarlink has no __module__ attribute:DeprecationWarning"
)

def test_pdf_dependencies_are_available():
    import fitz
    import pdfplumber

    assert fitz is not None
    assert pdfplumber is not None