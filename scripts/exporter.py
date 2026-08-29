"""
exporter.py
Assembles approved sections into a formatted Word document.
Uses a provided template if available, otherwise generates
a clean professional document from scratch.
"""

import os
import re
import logging
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────

STYLE_HEADING_1 = os.getenv("STYLE_HEADING_1", "Heading 1")
STYLE_HEADING_2 = os.getenv("STYLE_HEADING_2", "Heading 2")
STYLE_HEADING_3 = os.getenv("STYLE_HEADING_3", "Heading 3")
STYLE_HEADING_4 = os.getenv("STYLE_HEADING_4", "Heading 4")
STYLE_HEADING_5 = os.getenv("STYLE_HEADING_5", "Heading 5")
STYLE_NORMAL    = os.getenv("STYLE_NORMAL",    "Normal")

HEADING_STYLE_MAP = {
    1: STYLE_HEADING_1,
    2: STYLE_HEADING_2,
    3: STYLE_HEADING_3,
    4: STYLE_HEADING_4,
    5: STYLE_HEADING_5,
}


# ─────────────────────────────────────────────────────────
# DATA STRUCTURES
# ─────────────────────────────────────────────────────────

@dataclass
class Section:
    name:    str
    content: str
    level:   int


# ─────────────────────────────────────────────────────────
# TEMPLATE INSPECTION
# ─────────────────────────────────────────────────────────

def inspect_template_styles(template_path: Path) -> list[str]:
    """Lists all style names available in a template."""
    doc    = Document(template_path)
    styles = [s.name for s in doc.styles]
    logger.info(f"Styles found in template ({len(styles)}):")
    for style in sorted(styles):
        logger.info(f"  - {style}")
    return styles


def validate_styles(template_path: Path) -> dict[str, bool]:
    """
    Checks whether expected style names exist in the template.
    Logs a warning for any missing styles.
    """
    available = inspect_template_styles(template_path)
    expected  = {
        STYLE_HEADING_1: STYLE_HEADING_1 in available,
        STYLE_HEADING_2: STYLE_HEADING_2 in available,
        STYLE_HEADING_3: STYLE_HEADING_3 in available,
        STYLE_HEADING_4: STYLE_HEADING_4 in available,
        STYLE_HEADING_5: STYLE_HEADING_5 in available,
        STYLE_NORMAL:    STYLE_NORMAL    in available,
    }

    missing = [k for k, v in expected.items() if not v]
    if missing:
        logger.warning(
            f"Missing styles in template: {missing}. "
            f"Update STYLE_* variables in .env to match template style names."
        )
    else:
        logger.info("All expected styles found in template.")

    return expected


# ─────────────────────────────────────────────────────────
# DOCUMENT CREATION
# ─────────────────────────────────────────────────────────

def create_from_scratch() -> Document:
    """
    Creates a clean professional document with standard styles.
    Used when no template is provided.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    h1 = doc.styles[STYLE_HEADING_1]
    h1.font.size      = Pt(16)
    h1.font.bold      = True
    h1.font.color.rgb = RGBColor(0x00, 0x33, 0x99)

    h2 = doc.styles[STYLE_HEADING_2]
    h2.font.size      = Pt(13)
    h2.font.bold      = True
    h2.font.color.rgb = RGBColor(0x00, 0x55, 0x88)

    h3 = doc.styles[STYLE_HEADING_3]
    h3.font.size      = Pt(11)
    h3.font.bold      = True
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    normal = doc.styles[STYLE_NORMAL]
    normal.font.size = Pt(11)
    normal.font.name = "Calibri"

    logger.info("Created document from scratch with standard styles")
    return doc


def open_template(template_path: Path) -> Document:
    """Opens an existing Word template and validates its styles."""
    logger.info(f"Opening template: {template_path}")
    validate_styles(template_path)
    doc = Document(template_path)
    logger.info("Template loaded successfully")
    return doc


# ─────────────────────────────────────────────────────────
# CONTENT INSERTION
# ─────────────────────────────────────────────────────────

def clean_markdown(text: str) -> str:
    """Removes basic markdown markers before inserting into Word."""
    text = re.sub(r'^#{1,4}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    return text.strip()

def strip_redundant_section_heading(content: str, section_name: str) -> str:
    """
    Removes a leading Markdown heading when it repeats the section name.
    """
    lines = content.lstrip().splitlines()

    if not lines:
        return content

    first_line = clean_markdown(lines[0]).strip()
    normalized_name = clean_markdown(section_name).strip()

    if first_line.casefold() == normalized_name.casefold():
        return "\n".join(lines[1:]).lstrip()

    return content

def insert_markdown_table(doc: Document, markdown_table: str):
    """Converts a Markdown table into a Word table."""
    rows = [
        row for row in markdown_table.splitlines()
        if row.strip() and not set(
            row.replace("|", "").replace("-", "").strip()
        ) == set()
    ]

    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = max(len(row) for row in parsed)
    table    = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    doc.add_paragraph()
    logger.debug(f"Inserted table: {len(parsed)} rows x {num_cols} cols")


def insert_section(doc: Document, section: Section):
    """
    Inserts a single section into the document.
    Adds heading then content paragraphs or tables.
    """
    doc.add_heading(section.name, level=section.level)

    content = strip_redundant_section_heading(
        section.content,
        section.name,
    )

    paragraphs = [
        p.strip() for p in content.split("\n\n")
        if p.strip()
    ]

    for para_text in paragraphs:
        if para_text.startswith("|"):
            insert_markdown_table(doc, para_text)
        else:
            clean = clean_markdown(para_text)
            if clean:
                para = doc.add_paragraph(clean, style=STYLE_NORMAL)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()


# ─────────────────────────────────────────────────────────
# PUBLIC INTERFACE
# ─────────────────────────────────────────────────────────


def export_document(
    sections:      list[Section],
    project_name:  str,
    output_dir:    Path,
    template_path: Path | None = None,
) -> Path:
    """
    Main entry point. Assembles all sections into a Word document.

    Args:
        sections:      Ordered list of Section objects to export
        project_name:  Used for the output filename
        output_dir:    Where to save the .docx — passed from Streamlit session
        template_path: Optional path to a .docx template

    Returns:
        Path to the generated .docx file
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    # Open template or create from scratch
    if template_path and template_path.exists():
        doc = open_template(template_path)
    else:
        if template_path:
            logger.warning(
                f"Template not found at {template_path} — "
                f"generating document from scratch"
            )
        doc = create_from_scratch()

    # Insert sections
    logger.info(f"Inserting {len(sections)} sections into document")
    for section in sections:
        insert_section(doc, section)
        logger.info(f"  Inserted: {section.name}")

    # Save
    timestamp   = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name   = project_name.replace(" ", "_").lower()
    output_path = output_dir / f"{safe_name}_{timestamp}.docx"

    doc.save(output_path)
    logger.info(f"Document saved: {output_path.resolve()}")

    return output_path